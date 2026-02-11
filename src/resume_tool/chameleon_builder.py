import json
import os
from rich.console import Console
from rich.prompt import Prompt, Confirm
from rich.panel import Panel
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

console = Console()

# --- 1. LOAD DATABASE FROM FILE ---
def load_database():
    # Look for the file created by profile_builder.py
    if os.path.exists("master_profile.json"):
        with open("master_profile.json", "r") as f:
            return json.load(f)
    else:
        # Fallback to empty structure if file doesn't exist
        return {"summary": {}, "experience": [], "skills": {}}

def select_summary(data):
    console.print(Panel("[bold blue]Step 1: Choose a Summary[/bold blue]", style="blue"))
    
    if not data.get("summary"):
        console.print("[yellow]No summaries found in master profile.[/yellow]")
        return None
        
    options = list(data["summary"].keys())
    for idx, key in enumerate(options, 1):
        console.print(f"{idx}. {key}")
        
    choice = Prompt.ask("Select Summary", choices=[str(i) for i in range(1, len(options)+1)])
    selected_key = options[int(choice)-1]
    return data["summary"][selected_key]

def select_experience(data):
    console.print(Panel("[bold blue]Step 2: Filter Experience[/bold blue]", style="blue"))
    
    # Get all unique domains from experience
    all_domains = set()
    for entry in data.get("experience", []):
        for domain in entry.get("domains", []):
            all_domains.add(domain)
            
    if not all_domains:
        console.print("[yellow]No experience entries found with domains.[/yellow]")
        return []

    console.print(f"Available Domains: {', '.join(sorted(all_domains))}")
    target_role = Prompt.ask("Enter Target Role (e.g. QA, PRODUCT, TECH)").upper()
    
    selected_bullets = []
    console.print(f"\n[bold]Selecting bullets for {target_role}...[/bold]")
    
    for entry in data.get("experience", []):
        if target_role in entry.get("domains", []):
            console.print(f"[green]✓ Adding:[/green] {entry['text']}")
            selected_bullets.append(entry)
            
    return selected_bullets

def select_skills(data):
    console.print(Panel("[bold blue]Step 3: Select Skills[/bold blue]", style="blue"))
    
    if not data.get("skills"):
        console.print("[yellow]No skills found.[/yellow]")
        return {}
        
    selected_skills = {}
    for category, skills in data["skills"].items():
        if Confirm.ask(f"Include {category} skills?"):
            selected_skills[category] = skills
            
    return selected_skills

def generate_resume(summary, experience, skills):
    console.clear()
    console.print(Panel("[bold green]✨ GENERATING TARGETED RESUME ✨[/bold green]", style="green"))
    
    console.print("\n[bold]PROFESSIONAL SUMMARY[/bold]")
    console.print(summary)
    
    console.print("\n[bold]EXPERIENCE[/bold]")
    for entry in experience:
        console.print(f"• {entry['text']}")
        
    console.print("\n[bold]SKILLS[/bold]")
    for cat, skill_list in skills.items():
        console.print(f"[bold]{cat}:[/bold] {', '.join(skill_list)}")
        
    console.print("\n[dim]Resume generated successfully![/dim]")

def save_to_docx(filename, role_name, summary, bullets, skills):
    try:
        doc = Document()
        
        # Styling
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # Header
        doc.add_heading('J. Bateman', level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Baltimore, MD | (301) 213-0948 | j.batemansheppard@gmail.com | LinkedIn/JBateman').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('='*80)

        # Summary
        doc.add_heading('PROFESSIONAL SUMMARY', level=2)
        doc.add_paragraph(summary)

        # Experience
        doc.add_heading('PROFESSIONAL EXPERIENCE', level=2)
        p = doc.add_paragraph()
        runner = p.add_run('APKUDO LLC | Technical Delivery Analyst (Product Focus) | 2021 – Present')
        runner.bold = True
        
        # --- THE FIX IS HERE ---
        for b in bullets:
            # 1. Strip existing bullets/dashes/stars/question marks from the start (handle dict or string)
            text = b['text'] if isinstance(b, dict) else str(b)
            clean_text = text.lstrip("•-–*? ")
            
            # 2. Add the paragraph using Word's NATIVE bullet style
            doc.add_paragraph(clean_text, style='List Bullet')

        # Project Section
        doc.add_heading('Projects', level=2)
        p_proj = doc.add_paragraph()
        runner_proj = p_proj.add_run('FOCUSFINANCE | Product Owner & Developer')
        runner_proj.bold = True
        doc.add_paragraph('Designed and launched an Android budgeting application, defining the product roadmap based on user needs.')

        # Skills
        doc.add_heading('SKILLS & TOOLS', level=2)
        
        # Flatten skills dict to list if needed, or join values
        skill_text = ""
        if isinstance(skills, dict):
             all_skills = []
             for cat, s_list in skills.items():
                 all_skills.extend(s_list)
             skill_text = ' | '.join(all_skills)
        elif isinstance(skills, list):
             skill_text = ' | '.join(skills)
        else:
             skill_text = str(skills)

        doc.add_paragraph(skill_text)

        # Save
        doc.save(filename)
        print(f"\n✅ SUCCESS: Resume saved to {filename}")
        return True
    except Exception as e:
        print(f"\n❌ ERROR: Could not save file: {e}")
        return False

def main():
    data = load_database()
    
    if not data["summary"] and not data["experience"]:
        console.print("[red]Error: master_profile.json is empty or missing.[/red]")
        console.print("Run option 4 (Create Profile) first.")
        return

    summary = select_summary(data)
    experience = select_experience(data)
    skills = select_skills(data)
    
    generate_resume(summary, experience, skills)
    
    if Confirm.ask("\nSave this resume to DOCX?"):
        filename = Prompt.ask("Enter filename", default="Targeted_Resume.docx")
        # Extract just the skill names for the docx function
        save_to_docx(filename, "Targeted Role", summary, experience, skills)

if __name__ == "__main__":
    main()
