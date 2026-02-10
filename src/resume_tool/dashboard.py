import streamlit as st
import os
import sys
import io
import re
import random # Added for "Smart" variation
import fitz  # PyMuPDF
from docx import Document 
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF 
import nltk
try:
    import spacy
except (ImportError, OSError, Exception):
    spacy = None

# --- 1. CONFIG & STYLE ENGINE ---
st.set_page_config(
    page_title="Resume Command Center",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- CUSTOM CSS (The "Nike/LinkedIn" Look) ---
st.markdown("""
<style>
    /* 1. Main Background - LinkedIn Light Gray */
    .stApp {
        background-color: #f3f2ef;
    }
    
    /* 2. Card Styling - The "White Box" look */
    div[data-testid="stVerticalBlock"] > div > div {
        background-color: transparent;
    }
    .css-card {
        background-color: white;
        padding: 20px;
        border-radius: 10px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        margin-bottom: 20px;
    }
    
    /* 3. Typography - Nike Boldness */
    h1, h2, h3 {
        font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
        font-weight: 700;
        color: #111;
    }
    
    /* 4. Custom Button Styling - "Action" Blue */
    .stButton > button {
        background-color: #0a66c2; /* LinkedIn Blue */
        color: white;
        border-radius: 20px;
        font-weight: bold;
        border: none;
    }
    .stButton > button:hover {
        background-color: #004182;
    }

    /* 5. Success/Error Messages */
    .stAlert {
        border-radius: 8px;
    }
</style>
""", unsafe_allow_html=True)

current_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.append(current_dir)

# --- 2. LOAD LIBRARIES ---
@st.cache_resource
def load_nlp():
    try: nltk.data.find('tokenizers/punkt')
    except: nltk.download('punkt'); nltk.download('punkt_tab'); nltk.download('stopwords')
    if spacy:
        try: return spacy.load("en_core_web_sm")
        except: 
            try:
                from spacy.cli import download; download("en_core_web_sm"); return spacy.load("en_core_web_sm")
            except:
                return None
    return None

nlp = load_nlp()

try:
    from ats_auditor import get_analysis_data
    from career_data import CAREER_CLUSTERS
    from expert_tips import get_expert_advice, INDUSTRY_TRANSLATORS
except ImportError:
    st.error("⚠️ CRITICAL ERROR: Modules missing.")
    st.stop()

# --- 3. INTELLIGENT PARSER ---
def parse_contact_info(text):
    data = {"name": "", "email": "", "phone": "", "linkedin": ""}
    lines = text.split('\n')
    for line in lines[:5]:
        if line.strip() and len(line) < 50: data['name'] = line.strip(); break
    email = re.search(r'[\w\.-]+@[\w\.-]+', text)
    if email: data['email'] = email.group(0)
    phone = re.search(r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', text)
    if phone: data['phone'] = phone.group(0)
    if "linkedin" in text.lower(): data['linkedin'] = "LinkedIn Profile"
    return data

def parse_resume_sections(text):
    sections = {"summary": "", "skills": "", "experience": "", "education": ""}
    lines = text.split('\n')
    current_section = None
    buffer = []
    markers = {
        "summary": ["professional summary", "profile", "summary", "objective"],
        "skills": ["skills", "competencies", "technologies", "core competencies"],
        "experience": ["experience", "work history", "employment", "professional experience"],
        "education": ["education", "academic", "university", "college"]
    }
    for line in lines:
        clean_line = line.strip().lower()
        is_header = False
        for section, keywords in markers.items():
            if clean_line in keywords or (clean_line.endswith(":") and clean_line.strip(":") in keywords):
                if current_section: sections[current_section] = "\n".join(buffer).strip()
                current_section = section
                buffer = []
                is_header = True
                break
        if not is_header and current_section: buffer.append(line)
    if current_section: sections[current_section] = "\n".join(buffer).strip()
    return sections

# --- 4. EXPORT ENGINE ---
def reconstruct_resume_text(data):
    full_text = f"{data['name']}\n{data['email']} | {data['phone']}\n"
    if data['summary']: full_text += f"\nSummary:\n{data['summary']}\n"
    if data['skills']: full_text += f"\nSkills:\n{data['skills']}\n"
    if data['experience']: full_text += f"\nExperience:\n{data['experience']}\n"
    if data['education']: full_text += f"\nEducation:\n{data['education']}\n"
    return full_text

def create_docx(data):
    doc = Document()
    doc.add_heading(data['name'], 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"{data['email']} | {data['phone']} | {data['linkedin']}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    if data['summary']: doc.add_heading('SUMMARY', 1); doc.add_paragraph(data['summary'])
    if data['skills']: doc.add_heading('SKILLS', 1); doc.add_paragraph(data['skills'])
    if data['experience']: doc.add_heading('EXPERIENCE', 1); doc.add_paragraph(data['experience'])
    if data['education']: doc.add_heading('EDUCATION', 1); doc.add_paragraph(data['education'])
    buffer = io.BytesIO(); doc.save(buffer); buffer.seek(0); return buffer

def create_pdf(data):
    pdf = FPDF(); pdf.add_page(); pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, data['name'], ln=True, align='C')
    pdf.set_font("Arial", "", 10)
    pdf.cell(0, 5, f"{data['email']} | {data['phone']} | {data['linkedin']}", ln=True, align='C'); pdf.ln(5)
    def sect(t, b): 
        if b: 
            pdf.set_font("Arial","B",12); pdf.cell(0,8,t,ln=True,border='B'); pdf.ln(2)
            pdf.set_font("Arial","",10); pdf.multi_cell(0,5,b.encode('latin-1','replace').decode('latin-1')); pdf.ln(3)
    sect("SUMMARY", data['summary']); sect("SKILLS", data['skills']); sect("EXPERIENCE", data['experience']); sect("EDUCATION", data['education'])
    return pdf.output(dest='S').encode('latin-1')

def extract_pdf(f): 
    try: return "".join([p.get_text() for p in fitz.open(stream=f.read(), filetype="pdf")])
    except: return None

# --- 5. ROBUST AUTH SYSTEM ---
def check_authentication():
    # If already authenticated, return True
    if st.session_state.get("authenticated", False):
        return True
    
    # Login Form
    st.markdown("<h1 style='text-align: center;'>🔐 Resume Command Center</h1>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center;'>Please sign in to access your workspace.</p>", unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1,2,1])
    with col2:
        with st.container(border=True):
            password = st.text_input("Password", type="password")
            if st.button("Sign In", use_container_width=True):
                if password == "admin123": # <--- CHANGE THIS PASSWORD
                    st.session_state["authenticated"] = True
                    st.rerun()
                else:
                    st.error("Incorrect password.")
    return False

# --- 6. SESSION STATE ---
if 'resume_data' not in st.session_state:
    st.session_state['resume_data'] = {"name":"", "email":"", "phone":"", "linkedin":"", "summary":"", "skills":"", "experience":"", "education":""}
if 'resume_text' not in st.session_state: st.session_state['resume_text'] = ""
if 'jd_text' not in st.session_state: st.session_state['jd_text'] = ""
if 'missing_keywords' not in st.session_state: st.session_state['missing_keywords'] = []

# --- 7. MAIN APP FLOW ---
if check_authentication():
    
    # --- SIDEBAR (Global Inputs) ---
    with st.sidebar:
        st.title("🚀 Command Center")
        st.markdown("---")
        uploaded = st.file_uploader("Upload Resume (PDF)", type="pdf")
        if uploaded:
            text = extract_pdf(uploaded)
            if text:
                st.session_state['resume_text'] = text
                # Smart Populate (Preserve edits)
                if not st.session_state['resume_data']['name']:
                    st.session_state['resume_data'].update(parse_contact_info(text))
                    st.session_state['resume_data'].update(parse_resume_sections(text))
                    if not st.session_state['resume_data']['experience']: 
                        st.session_state['resume_data']['experience'] = text[300:] 
                st.success("✅ Profile Loaded")
        
        st.session_state['jd_text'] = st.text_area("Target Job Description", height=150, placeholder="Paste JD here to activate AI audit...")
        
        if st.button("Log Out"):
            st.session_state["authenticated"] = False
            st.rerun()

    # --- MAIN CONTENT ---
    
    # Header
    st.markdown(f"## Welcome back, {st.session_state['resume_data']['name'] or 'Candidate'}")
    
    if not st.session_state['resume_text']:
        st.info("👋 To begin, upload your PDF resume in the sidebar.")
    else:
        # TABS
        t1, t2, t3, t4 = st.tabs(["🔮 Career Pathfinder", "📊 Profile Audit", "🏗️ Smart Editor", "👨‍💻 Portfolio"])

        # TAB 1: PATHFINDER
        with t1:
            with st.container(border=True):
                st.subheader("Industry Compatibility")
                if st.button("🔍 Analyze My Fit"):
                    current_text = reconstruct_resume_text(st.session_state['resume_data'])
                    from career_data import CAREER_CLUSTERS
                    scores = []
                    for role, d in CAREER_CLUSTERS.items():
                        s = sum(1 for k in d['keywords'] if k in current_text.lower())
                        scores.append((role, s))
                    scores.sort(key=lambda x: x[1], reverse=True)
                    
                    # Display as Cards
                    cols = st.columns(3)
                    for i, (role, score) in enumerate(scores[:3]):
                        with cols[i]:
                            st.metric(label=role, value=f"{score} pts")
                            st.progress(min(score*10, 100)/100)

        # TAB 2: AUDITOR
        with t2:
            with st.container(border=True):
                st.subheader("ATS Intelligence")
                
                col_metric, col_advice = st.columns([1, 2])
                
                with col_metric:
                    if st.button("🚀 Run Scan", use_container_width=True):
                        current_draft_text = reconstruct_resume_text(st.session_state['resume_data'])
                        data = get_analysis_data(current_draft_text, st.session_state['jd_text'])
                        st.session_state['audit_score'] = data['match_score']
                        st.session_state['audit_feedback'] = data.get('feedback', [])
                        st.session_state['missing_keywords'] = data.get('missing_keywords', [])
                        st.session_state['common_keywords'] = data.get('common_keywords', [])

                    # Circular Style Metric
                    score = st.session_state.get('audit_score', 0)
                    st.markdown(f"<h1 style='text-align: center; font-size: 80px; color: #0a66c2;'>{score}%</h1>", unsafe_allow_html=True)
                    st.markdown("<p style='text-align: center;'>Match Score</p>", unsafe_allow_html=True)

                with col_advice:
                    if st.session_state.get('audit_feedback'):
                        st.write("### 🏥 Health Check")
                        for item in st.session_state['audit_feedback']: st.info(item)
                    elif st.session_state.get('missing_keywords'):
                        st.write("### ⚠️ Missing Keywords")
                        st.markdown("We found gaps in your profile. Go to the **Smart Editor** to add these:")
                        # Chip style display
                        st.markdown(" ".join([f"`{k}`" for k in st.session_state['missing_keywords'][:10]]))
                    else:
                        st.info("Run the scan to see insights.")

        # TAB 3: SMART EDITOR
        with t3:
            col_edit, col_prev = st.columns([1.2, 1])
            
            with col_edit:
                st.subheader("✏️ Workspace")
                
                # SMART SUGGESTION ENGINE (THE "SMARTER" PART)
                if st.session_state['missing_keywords']:
                    with st.expander("✨ AI Content Generator", expanded=True):
                        st.write("Click a missing keyword to generate a bullet point:")
                        kw_cols = st.columns(3)
                        for i, kw in enumerate(st.session_state['missing_keywords'][:6]):
                            if kw_cols[i%3].button(f"+ {kw}"):
                                # Mad Libs Generator
                                templates = [
                                    f"• Leveraged {kw} to optimize operational workflows, increasing efficiency by 20%.",
                                    f"• Spearheaded the integration of {kw}, resulting in improved data accuracy.",
                                    f"• Collaborated with cross-functional teams to implement {kw} best practices."
                                ]
                                st.session_state['resume_data']['experience'] += f"\n{random.choice(templates)}"
                                st.rerun()

                # PIVOT ENGINE
                with st.expander("⚡ Industry Translator", expanded=False):
                    target_industry = st.selectbox("Pivot to:", ["Select...", "Sports Marketing", "Healthcare", "Tech", "Finance"])
                    if target_industry != "Select..." and st.button(f"Translate"):
                        from expert_tips import INDUSTRY_TRANSLATORS
                        dictionary = INDUSTRY_TRANSLATORS.get(target_industry, {})
                        exp_text = st.session_state['resume_data']['experience']
                        for old, new in dictionary.items():
                            pattern = re.compile(r'\b' + re.escape(old) + r'\b', re.IGNORECASE)
                            exp_text = pattern.sub(new.upper(), exp_text)
                        st.session_state['resume_data']['experience'] = exp_text
                        st.success("Translation Complete")

                # EDIT FIELDS
                st.text_input("Full Name", key="name_input", value=st.session_state['resume_data']['name'], 
                              on_change=lambda: st.session_state['resume_data'].update({'name': st.session_state.name_input}))
                
                st.text_area("Professional Summary", key="summary_input", value=st.session_state['resume_data']['summary'],
                             on_change=lambda: st.session_state['resume_data'].update({'summary': st.session_state.summary_input}))
                
                st.text_area("Experience (Bullets)", key="exp_input", value=st.session_state['resume_data']['experience'], height=300,
                             on_change=lambda: st.session_state['resume_data'].update({'experience': st.session_state.exp_input}))
                
                st.text_area("Skills", key="skills_input", value=st.session_state['resume_data']['skills'],
                             on_change=lambda: st.session_state['resume_data'].update({'skills': st.session_state.skills_input}))

            with col_prev:
                st.subheader("📄 Live Preview")
                # Visual Paper Container
                st.markdown("""
                <div style="background-color: white; color: black; padding: 40px; border: 1px solid #ddd; box-shadow: 2px 2px 10px rgba(0,0,0,0.1); min-height: 600px;">
                """, unsafe_allow_html=True)
                
                d = st.session_state['resume_data']
                st.markdown(f"<h2 style='text-align:center; color:black;'>{d['name']}</h2>", unsafe_allow_html=True)
                st.markdown(f"<p style='text-align:center; color:#555;'>{d['email']} | {d['phone']} | {d['linkedin']}</p><hr>", unsafe_allow_html=True)
                
                if d['summary']: st.markdown(f"**SUMMARY**<br>{d['summary']}", unsafe_allow_html=True)
                if d['skills']: st.markdown(f"<br>**SKILLS**<br>{d['skills']}", unsafe_allow_html=True)
                if d['experience']: st.markdown(f"<br>**EXPERIENCE**<br>{d['experience'].replace(chr(10), '<br>')}", unsafe_allow_html=True)
                
                st.markdown("</div>", unsafe_allow_html=True)
                
                c1, c2 = st.columns(2)
                with c1: st.download_button("⬇️ PDF", create_pdf(st.session_state['resume_data']), "Resume.pdf", use_container_width=True)
                with c2: st.download_button("⬇️ DOCX", create_docx(st.session_state['resume_data']), "Resume.docx", use_container_width=True)

        # TAB 4: PORTFOLIO
        with t4:
            st.header("Under the Hood")
            st.markdown("### Architecture Diagram")
            st.markdown("```mermaid\ngraph LR\nA[User] --> B[Streamlit UI]; B --> C[Smart Parser]; C --> D{Logic Engine}; D -->|No JD| E[Health Check]; D -->|With JD| F[Vector Analysis]; D -->|Edit| G[Draft State];\n```")
